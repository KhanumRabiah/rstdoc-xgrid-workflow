#!/usr/bin/env python
# encoding: utf-8

from .reimg import main as reimg
from .reflow import main as reflow
from .untable import main as untable
from .listtable import main as listtable
from .dcx import example_rest_tree, rindices
from pathlib import Path
import time
import re
from glob import glob
import shutil
import os
import pypandoc
from zipfile import ZipFile
from rstdoc import __version__
import subprocess

# Added imports for code block detection
try:
    from docx import Document
    from docx.shared import Inches
    from docx.oxml.ns import qn
    from pandoc.types import RawBlock # This import might be needed for the filter
except ImportError:
    Document = None
    print("Warning: python-docx not installed. Code block detection will be disabled.")


"""
.. _`rstfromdocx`:

rstfromdocx
===========

| rstfromdocx: shell command
| fromdocx: rstdoc module

Convert DOCX to RST in a subfolder of current dir, named after the DOCX file.
It also creates ``conf.py``, ``index.py`` and ``Makefile``
and copies ``dcx.py`` into the folder.

See |rstdcx| for format conventions for the RST.

There are options to post-process through::

    --listtable (--join can be provided)
    --untable
    --reflow (--sentence True,  --join 0)
    --reimg

``rstfromdocx -lurg`` combines all of these.

To convert more DOCX documents into the same
RST documentation folder, proceed like this:

- rename/copy the original DOCX to the name you want for the ``.rst`` file
- run ``rstfromdocx -lurg doc1.docx``; instead of -lurg use your own options
- check the output in the ``doc1`` subfolder
- repeat the previous 2 steps with the next DOCX files
- create a new folder, e.g. ``doc``
- merge all other folders into that new folder

``fromdocx.docx_rst_5`` creates 5 different rst files with different postprocessing.

See |rstreflow| for an alternative proceeding.


"""



'''
API
---


.. code-block:: py

   import rstdoc.fromdocx as fromdocx

'''

# Global variables to track state across multiple file conversions
image_counter = 1
image_path_map = {} # Maps original image names to new sequential names

# ==================== CODE BLOCK DETECTION FUNCTIONS ====================

def detect_code_blocks(key, value, format, meta):
    """
    Pandoc filter function to detect preformatted text in Word documents
    and convert it to proper code-block directives in RST.
    
    This function processes Para elements that contain preformatted text
    (typically indented code blocks from Word) and converts them to
    RST code-block directives.
    """
    
    if key == "Para":  # Word preformatted text might be wrapped in Para
        text = value[0].get("c", "")
        if isinstance(text, str) and text.startswith(" "):  # Detect indented text
            return RawBlock("rst", ".. code-block::\n\n    " + text.replace("\n", "\n    "))


def detect_code_block_style(paragraph):
    """
    Detect if a paragraph represents a code block based on various indicators.
    
    Args:
        paragraph: python-docx paragraph object
        
    Returns:
        tuple: (is_code_block, detected_language)
    """
    # Check paragraph style name
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    
    # Common code block style names in Word documents
    code_style_patterns = [
        'code', 'source', 'literal', 'monospace', 'courier',
        'verbatim', 'preformatted', 'console', 'terminal'
    ]
    
    # Check if style name contains code-related keywords
    for pattern in code_style_patterns:
        if pattern in style_name:
            return True, detect_language_from_content(paragraph.text)
    
    # Check font characteristics
    if paragraph.runs:
        first_run = paragraph.runs[0]
        font = first_run.font
        
        # Check for monospace fonts
        monospace_fonts = [
            'courier', 'consolas', 'monaco', 'menlo', 'source code pro',
            'fira code', 'ubuntu mono', 'dejavu sans mono', 'liberation mono'
        ]
        
        font_name = font.name.lower() if font.name else ""
        for mono_font in monospace_fonts:
            if mono_font in font_name:
                return True, detect_language_from_content(paragraph.text)
    
    # Check paragraph formatting
    paragraph_format = paragraph.paragraph_format
    if paragraph_format:
        # Check for specific indentation patterns
        if (paragraph_format.left_indent and
            paragraph_format.left_indent.pt > 36):  # Significant indentation
            # Additional check: does it contain code-like patterns?
            if contains_code_patterns(paragraph.text):
                return True, detect_language_from_content(paragraph.text)
    
    # Check content patterns
    if contains_code_patterns(paragraph.text):
        return True, detect_language_from_content(paragraph.text)
    
    return False, None


def contains_code_patterns(text):
    """
    Check if text contains patterns typical of code.
    
    Args:
        text: string to analyze
        
    Returns:
        bool: True if text appears to contain code
    """
    if not text.strip():
        return False
    
    # Code-like patterns
    code_indicators = [
        r'\b(def|class|import|from|if|else|elif|for|while|try|except|return)\b',  # Python keywords
        r'\b(function|var|let|const|if|else|for|while|return)\b',  # JavaScript keywords
        r'\b(public|private|static|class|interface|import|package)\b',  # Java keywords
        r'\b(SELECT|FROM|WHERE|INSERT|UPDATE|DELETE|CREATE|TABLE)\b',  # SQL keywords
        r'[{}();]',  # Common programming punctuation
        r'^\s*[#//]',  # Comments
        r'^\s*<[^>]+>',  # HTML/XML tags
        r'\$\w+',  # Shell variables
        r'^\s*\w+\s*=',  # Variable assignments
        r'^\s*\.',  # Method calls or relative paths
        r'^\s*@\w+',  # Decorators or annotations
    ]
    
    pattern_count = 0
    for pattern in code_indicators:
        if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
            pattern_count += 1
            if pattern_count >= 2:  # Multiple indicators suggest code
                return True
    
    # Check for high density of special characters
    special_chars = sum(1 for c in text if c in '{}[]();,.:=<>+-*/%&|^~')
    if len(text) > 0 and special_chars / len(text) > 0.1:
        return True
    
    return False


def detect_language_from_content(text):
    """
    Attempt to detect programming language from code content.
    
    Args:
        text: code content string
        
    Returns:
        str: detected language or 'text' as default
    """
    if not text.strip():
        return 'text'
    
    # Language detection patterns
    language_patterns = {
        'python': [
            r'\bdef\s+\w+\s*\(',
            r'\bimport\s+\w+',
            r'\bfrom\s+\w+\s+import',
            r'^\s*#.*$',
            r'\bprint\s*\(',
            r'\b__\w+__\b'
        ],
        'javascript': [
            r'\bfunction\s+\w+\s*\(',
            r'\b(var|let|const)\s+\w+',
            r'\bconsole\.log\s*\(',
            r'^\s*//.*$',
            r'=>',
            r'\$\{.*\}'
        ],
        'java': [
            r'\bpublic\s+class\s+\w+',
            r'\bpublic\s+static\s+void\s+main',
            r'\bSystem\.out\.println',
            r'\bpublic\s+\w+\s+\w+\s*\(',
            r'\bimport\s+[\w.]+;'
        ],
        'sql': [
            r'\bSELECT\s+.*\bFROM\b',
            r'\bINSERT\s+INTO\b',
            r'\bUPDATE\s+.*\bSET\b',
            r'\bCREATE\s+TABLE\b',
            r'\bDELETE\s+FROM\b'
        ],
        'html': [
            r'<html[^>]*>',
            r'<head[^>]*>',
            r'<body[^>]*>',
            r'<div[^>]*>',
            r'<!DOCTYPE\s+html>'
        ],
        'css': [
            r'\w+\s*\{[^}]*\}',
            r'@media\s*\([^)]*\)',
            r'#\w+\s*\{',
            r'\.\w+\s*\{'
        ],
        'bash': [
            r'^#!/bin/(bash|sh)',
            r'\$\w+',
            r'\becho\s+',
            r'\bif\s+\[.*\];\s*then',
            r'\|\s*\w+'
        ],
        'json': [
            r'^\s*\{',
            r'"\w+"\s*:',
            r'^\s*\[',
            r'}\s*,?\s*$'
        ],
        'xml': [
            r'<\?xml\s+version',
            r'<\w+[^>]*>.*</\w+>',
            r'xmlns\s*='
        ]
    }
    
    # Count matches for each language
    language_scores = {}
    for language, patterns in language_patterns.items():
        score = 0
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE | re.MULTILINE):
                score += 1
        if score > 0:
            language_scores[language] = score
    
    # Return the language with the highest score
    if language_scores:
        return max(language_scores, key=language_scores.get)
    
    return 'text'


def convert_code_block_to_rst(text, language=None):
    """
    Convert code text to reStructuredText format.
    
    Args:
        text: code content string
        language: detected or specified language
        
    Returns:
        str: formatted reStructuredText code block
    """
    if not language:
        language = 'text'
    
    # Preserve original indentation and formatting
    lines = text.split('\n')
    
    # Build the rst code block
    rst_lines = [f'.. code-block:: {language}', '']
    
    # Add each line with proper indentation (3 spaces for rst)
    for line in lines:
        if line.strip():  # Non-empty lines
            rst_lines.append(f'   {line}')
        else:  # Empty lines
            rst_lines.append('')
    
    # Add extra newline at the end
    rst_lines.append('')
    
    return '\n'.join(rst_lines)


def is_consecutive_code_block(paragraphs, start_index):
    """
    Check if consecutive paragraphs form a single code block.
    
    Args:
        paragraphs: list of paragraph objects
        start_index: starting index to check from
        
    Returns:
        tuple: (end_index, combined_text, language)
    """
    if start_index >= len(paragraphs):
        return start_index, "", "text"
    
    combined_lines = []
    current_language = None
    end_index = start_index
    
    # Check consecutive paragraphs
    for i in range(start_index, len(paragraphs)):
        paragraph = paragraphs[i]
        is_code, lang = detect_code_block_style(paragraph)
        
        if is_code:
            combined_lines.append(paragraph.text)
            if not current_language and lang:
                current_language = lang
            end_index = i
        else:
            # Stop if we hit a non-code paragraph
            break
    
    combined_text = '\n'.join(combined_lines)
    return end_index, combined_text, current_language or 'text'


def enhance_rst_with_code_blocks(rst_text, docx_path):
    """
    Enhanced RST processing that detects and converts code blocks from the original DOCX.
    This function processes the DOCX file directly to detect code blocks and enhances
    the RST output with proper code-block directives.
    
    Args:
        rst_text: Original RST text from pandoc conversion
        docx_path: Path to the original DOCX file
        
    Returns:
        str: Enhanced RST text with code blocks converted to proper directives
    """
    if Document is None:
        return rst_text
    try:
        # Load the original DOCX document
        document = Document(docx_path)
        
        # Process paragraphs to find code blocks
        code_blocks_info = []
        i = 0
        
        while i < len(document.paragraphs):
            paragraph = document.paragraphs[i]
            is_code, detected_language = detect_code_block_style(paragraph)
            
            if is_code:
                # Check for consecutive code blocks
                end_index, combined_text, language = is_consecutive_code_block(document.paragraphs, i)
                
                if combined_text.strip():
                    code_blocks_info.append({
                        'text': combined_text,
                        'language': language,
                        'start_index': i,
                        'end_index': end_index
                    })
                
                # Skip the processed paragraphs
                i = end_index + 1
            else:
                i += 1
        
        # If no code blocks found, return original RST
        if not code_blocks_info:
            return rst_text
        
        # Enhance the RST text by replacing detected code blocks
        enhanced_rst = rst_text
        
        for code_info in code_blocks_info:
            # Find the code text in the RST and replace it with proper code-block directive
            code_text = code_info['text'].strip()
            language = code_info['language']
            
            # Create the proper RST code block
            rst_code_block = convert_code_block_to_rst(code_text, language)
            
            # Try to find and replace the code in the RST text
            # This is a heuristic approach - we look for the code text and replace it
            if code_text in enhanced_rst:
                enhanced_rst = enhanced_rst.replace(code_text, rst_code_block, 1)
            else:
                # If exact match fails, try to find similar patterns
                lines = code_text.split('\n')
                if lines:
                    first_line = lines[0].strip()
                    if first_line and first_line in enhanced_rst:
                        # Find the block and replace it
                        rst_lines = enhanced_rst.split('\n')
                        for j, rst_line in enumerate(rst_lines):
                            if first_line in rst_line:
                                # Try to identify the full code block in RST
                                block_start = j
                                block_end = j
                                
                                # Look for the end of the code block
                                for k in range(j + 1, len(rst_lines)):
                                    if rst_lines[k].strip() == '' or not rst_lines[k].startswith('   '):
                                        break
                                    block_end = k
                                
                                # Replace the identified block
                                rst_block_lines = rst_code_block.split('\n')
                                rst_lines[block_start:block_end + 1] = rst_block_lines
                                enhanced_rst = '\n'.join(rst_lines)
                                break
        
        return enhanced_rst
        
    except Exception as e:
        print(f"Warning: Could not enhance code blocks from DOCX: {e}")
        return rst_text

# NEW function to fix image references after conversion
def fix_image_references(rst_text):
    """
    Replaces original image filenames in RST text with the new sequential names.
    """
    global image_path_map
    processed_text = rst_text
    for original_name, new_name in image_path_map.items():
        processed_text = processed_text.replace(f"media/{original_name}", f"media/{new_name}")
    return processed_text

def _mkdir(fn):
    os.makedirs(fn, exist_ok=True)


def _rstname(fn):
    return os.path.splitext(os.path.split(fn)[1])[0]


def _prj_name(fn):
    m = re.match(r'[\s\d\W]*([^\s\W]*).*', _rstname(fn))
    return m.group(1).strip('_').replace(' ', '')

def _fldrhere(n):
    return os.getcwd()

# MODIFIED to populate the global image_path_map
def extract_media(adocx):
    '''
    Extract media files, rename them sequentially, and map old names to new names.
    '''
    global image_counter, image_path_map
    try:
        with ZipFile(adocx) as zf:
            media_dir = os.path.join(os.getcwd(), "media")
            _mkdir(media_dir)
            media_files = [x for x in zf.infolist() if 'media/' in x.filename and not x.filename.endswith('/')]
            
            for m in media_files:
                original_basename = os.path.basename(m.filename)
                file_ext = os.path.splitext(original_basename)[1]
                new_filename = f"image_{image_counter}{file_ext}"
                new_filepath = os.path.join(media_dir, new_filename)
                
                while os.path.exists(new_filepath):
                    image_counter += 1
                    new_filename = f"image_{image_counter}{file_ext}"
                    new_filepath = os.path.join(media_dir, new_filename)
                
                with zf.open(m) as source, open(new_filepath, "wb") as target:
                    shutil.copyfileobj(source, target)
                
                image_path_map[original_basename] = new_filename
                image_counter += 1
    except Exception as e:
        print(f"Warning: Could not extract media from {adocx}: {e}")

def _docxrst(adocx):
    _, frm = os.path.split(adocx)
    fnrst = os.path.splitext(frm)[0]
    return os.path.join(os.getcwd(), fnrst + '.rst')

def _write_confpy(adocx):
    cpfn = os.path.join(os.getcwd(), "conf.py")
    if os.path.exists(cpfn):
        return
    confpy = re.split(r'\s*.\s*Makefile', example_rest_tree.split('conf.py')[1])[0]
    pn = _prj_name(adocx)
    confpy = confpy.replace('docxsample', pn).replace('2017', time.strftime('%Y'))
    lns = confpy.splitlines(True)
    s = re.search(r'\w', lns[1]).span(0)[0]
    confpy = ''.join([l[s:] for l in lns])
    with open(cpfn, 'w', encoding='utf-8') as f:
        f.write(confpy)

def _write_index(adocx):
    ifn = os.path.join(os.getcwd(), "index.rst")
    rst_name = _rstname(adocx)
    prjname = _prj_name(adocx)
    hp = '=' * len(prjname)

    if os.path.exists(ifn):
        with open(ifn, 'r', encoding='utf-8') as f:
            lns = f.readlines()
    else:
        lns = [f"{hp}\n", f"{prjname}\n", f"{hp}\n\n", ".. toctree::\n", "   :maxdepth: 2\n\n"]
    
    rst_line_to_add = f'   {rst_name}\n'
    if rst_line_to_add not in lns:
        try:
            itoc_list = [i for i, line in enumerate(lns) if 'toctree' in line]
            if not itoc_list: raise IndexError
            itoc = itoc_list[0]
            insert_pos = itoc + 1
            while insert_pos < len(lns) and (lns[insert_pos].strip().startswith(':') or not lns[insert_pos].strip()):
                insert_pos += 1
            lns.insert(insert_pos, rst_line_to_add)
        except IndexError:
             lns.extend(['\n.. toctree::\n', '   :maxdepth: 2\n\n', rst_line_to_add])

    with open(ifn, 'w', encoding='utf-8') as f:
        f.writelines(lns)

def _write_makefile(adocx):
    mffn = os.path.join(os.getcwd(), "Makefile")
    if os.path.exists(mffn):
        return
    mf = re.split(r'\s\s+__code__', re.split('\s\sMakefile', example_rest_tree)[1])[0]
    lns = mf.splitlines(True)
    s = re.search(r'\w', lns[1]).span(0)[0]
    lns = [l[s:] for l in lns]
    with open(mffn, 'w', encoding='utf-8') as f:
        f.writelines(lns)

def process_rst_admonitions(rst_text):
    """
    Correctly converts multi-line admonitions into properly indented RST directives.
    """
    directive_mappings = {"note": "note", "see also": "seealso", "attention": "attention", "caution": "caution", "error": "error", "danger": "danger", "hint": "hint", "tip": "tip", "important": "important", "warning": "warning"}
    keywords_pattern = "|".join(re.escape(key) for key in directive_mappings.keys())
    pattern = re.compile(rf"^\s*[\*_]*\s*({keywords_pattern})\s*[\*_]*\s*:\s*(.*)", re.IGNORECASE)
    paragraphs = re.split(r'\n\s*\n', rst_text)
    processed_paragraphs = []
    for paragraph in paragraphs:
        if not paragraph.strip():
            processed_paragraphs.append(paragraph)
            continue
        lines = paragraph.split('\n')
        match = pattern.match(lines[0].strip())
        if match:
            keyword = match.group(1).lower().strip()
            rest_of_first_line = match.group(2).strip()
            directive = directive_mappings.get(keyword, "note")
            new_lines = [f".. {directive}::" + (f" {rest_of_first_line}" if rest_of_first_line else "")]
            new_lines.extend(f"   {line.rstrip()}" if line.strip() else "   " for line in lines[1:])
            processed_paragraphs.append('\n'.join(new_lines))
        else:
            processed_paragraphs.append(paragraph)
    return '\n\n'.join(processed_paragraphs)

## CORE LOGIC AND MAIN FUNCTIONS

def _process_single_file(args):
    '''Processes a single docx file based on provided arguments.'''
    adocx = args['docx']
    extract_media(adocx)
    fnrst = _docxrst(adocx)
    
    rst = pypandoc.convert_file(adocx, 'rst', 'docx')
    processed_rst = process_rst_admonitions(rst)
    enhanced_rst = enhance_rst_with_code_blocks(processed_rst, adocx)
    final_rst = fix_image_references(enhanced_rst) # Apply image path fix

    with open(fnrst, 'w', encoding='utf-8', newline='\n') as f:
        f.write('.. vim: syntax=rst\n\n')
        f.write(final_rst)
        
    _write_confpy(adocx)
    _write_index(adocx)
    _write_makefile(adocx)
    
    for processor in ['listtable', 'untable', 'reflow', 'reimg']:
        if args.get(processor):
            import argparse
            proc_args = args.copy()
            proc_args.update({'in_place': True, 'sentence': True, 'rstfile': [argparse.FileType('r', encoding='utf-8')(fnrst)]})
            if processor == 'reflow':
                proc_args['join'] = '0'
            eval(processor)(**proc_args)
    return fnrst

def batch_convert(directory, args):
    '''Converts all DOCX files in a directory and then runs sphinx-build.'''
    print(f"Starting batch conversion in directory: {directory}")
    for filename in sorted(os.listdir(directory)):
        if filename.lower().endswith(".docx"):
            input_file = os.path.join(directory, filename)
            print(f"Processing: {os.path.basename(input_file)}")
            file_args = args.copy()
            file_args['docx'] = input_file
            _process_single_file(file_args)
    print("\nBatch conversion complete!")

    build_dir = "build"
    source_dir = "." 
    print(f"\nBuilding Sphinx documentation...")
    try:
        result = subprocess.run(["sphinx-build", "-b", "html", source_dir, build_dir], check=True, capture_output=True, text=True)
        print(f"Build successful! HTML files are in: {os.path.join(os.getcwd(), build_dir)}")
    except FileNotFoundError:
        print("\nError: 'sphinx-build' command not found. Please ensure Sphinx is installed ('pip install -U Sphinx').")
    except subprocess.CalledProcessError as e:
        print(f"\nAn error occurred during sphinx-build:\n{e.stderr}")

def main(**args):
    '''
    Corresponds to the rstfromdocx shell command. Can process a single file or a directory.
    '''
    import argparse
    
    if not args:
        parser = argparse.ArgumentParser(description='''Convert a DOCX file or a directory of DOCX files to RST.''')
        parser.add_argument('--version', action='version', version = __version__)
        parser.add_argument('path', action='store', help='Path to a DOCX file or a directory containing DOCX files')
        parser.add_argument('-o', '--output-dir', action='store', default=None, help='Directory to save all output files.')
        parser.add_argument('-l', '--listtable', action='store_true', default=False, help='Postprocess through rstlisttable')
        parser.add_argument('-u', '--untable', action='store_true', default=False, help='Postprocess through rstuntable')
        parser.add_argument('-r', '--reflow', action='store_true', default=False, help='Postprocess through rstreflow')
        parser.add_argument('-g', '--reimg', action='store_true', default=False, help='Postprocess through rstreimg')
        parser.add_argument('-j', '--join', action='store', default='012', help='Join method for listtable, e.g., 002.')
        args = parser.parse_args().__dict__

    original_cwd = os.getcwd()
    path = os.path.abspath(args['path'])
    args['path'] = path

    output_dir = args.get('output_dir')
    if not output_dir:
        base_name = os.path.basename(path)
        if os.path.isdir(path):
            output_dir = base_name + "_docs"
        else:
            output_dir = os.path.splitext(base_name)[0] + "_docs"
    
    output_dir = os.path.join(original_cwd, output_dir)
    os.makedirs(output_dir, exist_ok=True)
    os.chdir(output_dir)
    print(f"Output will be saved in: {output_dir}")

    try:
        if os.path.isdir(path):
            batch_convert(path, args)
        elif os.path.isfile(path) and path.lower().endswith('.docx'):
            args['docx'] = args.pop('path')
            fnrst = _process_single_file(args)
            print(f"Conversion complete. Generated file: {fnrst}")
        else:
            print(f"Error: Path '{path}' is not a valid DOCX file or directory.")
    finally:
        os.chdir(original_cwd) # Always return to the original directory


def docx_rst_5(docx ,rename ,sentence=True):
    '''
    Creates 5 rst files in a sub-directory named after the 'rename' parameter.
    '''
    original_cwd = os.getcwd()
    output_dir = os.path.join(original_cwd, rename)
    _mkdir(output_dir)
    
    try:
        temp_docx = os.path.join(output_dir, os.path.basename(rename) + ".docx")
        shutil.copy2(docx, temp_docx)
        os.chdir(output_dir)
        
        # This function creates its own internal logic, we adapt to it
        fnrst = _process_single_file({'docx': temp_docx}) 
        r, _ = os.path.splitext(os.path.basename(fnrst))
        
        shutil.copy2(fnrst, r + '_l.rst')
        listtable(rstfile=r + '_l.rst', in_place=True)
        shutil.copy2(r + '_l.rst', r + '_u.rst')
        untable(rstfile=r + '_u.rst', in_place=True)
        shutil.copy2(r + '_u.rst', r + '_r.rst')
        reflow(rstfile=r + '_r.rst', in_place=True, sentence=sentence)
        shutil.copy2(r + '_r.rst', r + '_g.rst')
        reimg(rstfile=r + '_g.rst', in_place=True)
    finally:
        os.chdir(original_cwd)


if __name__ == '__main__':
    main()

# vim: ts=4 sw=4 sts=4 et noai nocin nosi